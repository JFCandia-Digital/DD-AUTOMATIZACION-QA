#!/usr/bin/env python3
"""Genera Informe QA Nodo PISEE — formato estructurado para Asana/Word."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/api-v3.5/Informe-QA-Nodo-PISEE-DocDigital-02-07-2026.docx"


def shade(cell, color):
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(s)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("INFORME QA — Integración Nodo PISEE ↔ DocDigital API v3.5\n")
    r.bold = True
    r.font.size = Pt(16)
    r2 = t.add_run("Pruebas subtareas error + flujo integración | Ambiente TEST\n")
    r2.font.size = Pt(12)
    r3 = t.add_run("Fecha informe: 02-07-2026 | Ejecutor: Juan Francisco Candia (JFC)\n")
    r3.italic = True

    doc.add_paragraph()

    # Estado global
    p = doc.add_paragraph()
    p.add_run("ESTADO GLOBAL: ").bold = True
    p.add_run(
        "Flujo integración PASS. Subtareas de error ejecutadas (8/8). "
        "Hallazgo crítico en nodo: no propaga códigos HTTP de error de DocDigital. "
        "Pendiente validación negocio en escenarios cross-entidad GET (H-03)."
    )

    doc.add_paragraph()

    # Tabla contexto
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Proyecto", "Integración Nodo PISEE consumidor DocDigitalv3"),
        ("Ambiente", "TEST — https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Nodo", "NodoV2 v2.4.0 | PE-TEST-00001 | localhost:8085/ddv3/*"),
        ("Token flujo OK", "PDI (entidad 66)"),
        ("Token cross-entidad", "ARMADA (entidad 210)"),
        ("Evidencias", "Documento_Evidencias_Pruebas NODO PISSE.docx (EV-01 a EV-19)"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()

    # 1 Hallazgos
    heading(doc, "1. Hallazgos y observaciones (prioridad)", 1)

    hallazgos = [
        (
            "H-01 — CRÍTICO",
            "El nodo PISEE no propaga los códigos HTTP de error de DocDigital API v3.5.",
            [
                "Ante 401 (sin token, token inválido, token truncado), 403 (acuse cross-entidad), "
                "404 (recurso/ruta inexistente) y 400 (validación) en API directa, el nodo responde "
                "500 Internal Server Error con body vacío.",
                "Solo queda cabecera Id_traza para escalamiento a Interoperabilidad/PISEE.",
                "Ejemplos Id_traza: 01KWHQW4… (401), 01KWJDRY5… (400), 01KWJDDDR… (403 acuse), "
                "01KWHR3T… (404).",
                "Impacto: integradores que consumen solo vía nodo no pueden interpretar el error real.",
            ],
        ),
        (
            "H-02 — MEDIO",
            "Ruta incorrecta en guías antiguas: GET /entidades/token.",
            [
                "Directo TEST: 404 — No static resource entidades/token.",
                "Nodo: 500 body vacío.",
                "Ruta correcta Swagger: GET /entidades/entidad-token → 200 directo y nodo.",
            ],
        ),
        (
            "H-03 — MEDIO",
            "Comportamiento cross-entidad inconsistente en GET vs PUT acuse.",
            [
                "Token ARMADA sobre com. PDI 89757 (GET detalle): directo 200 y nodo 200 — "
                "no reproduce 401 (subtarea #4) ni 403 detalle (subtarea #5).",
                "Token ARMADA sobre acuse com. PDI 89758 / tarea 618993 (PUT): directo 403 "
                "Acceso denegado — subtarea #6 PASS.",
                "Nodo en acuse 403: 500 vacío (H-01).",
                "Recomendación: confirmar con negocio/backend si GET cross-entidad en TEST es esperado.",
            ],
        ),
        (
            "H-04 — INFO",
            "Desalineación tarjeta vs catálogo.",
            [
                "Tarjeta Asana indicaba DEMO; catálogo nodo apuntaba a TEST.",
                "Equipo confirmó que TEST alcanza para cierre de pruebas.",
            ],
        ),
    ]

    for hid, title, items in hallazgos:
        p = doc.add_paragraph()
        p.add_run(f"{hid}: ").bold = True
        p.add_run(title)
        for item in items:
            bullet(doc, item)

    doc.add_page_break()

    # 2 Flujo PASS
    heading(doc, "2. Flujo integración — Camino feliz (PASS)", 1)

    bullet(doc, "NodoV2 servicios DocDigitalv3 — Endpoint TEST verificado.", "EV-01: ")
    bullet(doc, "POST /ddv3/oauth/token — 200, token PDI.", "EV-02: ")
    bullet(doc, "POST test-recepcion — 200, com. 89757 (folio QA-REC-JFC-NODO-003).", "EV-03: ")
    bullet(doc, "GET pendientes-recepcion — 200, tarea 618990.", "EV-04: ")
    bullet(doc, "PUT acuse-recibo PDI — 200, completada true.", "EV-05: ")
    bullet(doc, "Correo DocDigital recibido.", "EV-06: ")
    bullet(doc, "GET /comunicaciones/89757 — 200 directo y 200 nodo (equivalentes).", "EV-07: ")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Conclusión flujo OK: ").bold = True
    p.add_run(
        "El nodo transporta correctamente las respuestas 200 del happy path. "
        "La integración funcional recepción/acuse/consulta está validada en TEST."
    )

    doc.add_page_break()

    # 3 Subtareas
    heading(doc, "3. Resultado por subtarea Asana (error vía Nodo DD)", 1)

    doc.add_paragraph(
        "Cada subtarea se probó comparando API directa TEST vs nodo local (localhost:8085/ddv3). "
        "Criterio: la API directa debe devolver el código HTTP documentado; se registra si el nodo "
        "propaga o transforma la respuesta."
    )

    subtareas = [
        (
            "3.1 — Test error 404: endpoint inexistente o nombre incorrecto",
            "EV-08, EV-09, EV-10, EV-11",
            [
                ("999999 inexistente", "GET /comunicaciones/999999 + PDI", "404", "404 Registro no encontrado", "500 vacío", "01KWHR3T…", "OK + H-01"),
                ("Ruta incorrecta", "GET /entidades/token + PDI", "404", "404 No static resource", "500 vacío", "01KWJ8A9…", "OK + H-02"),
                ("Ruta correcta (contraste)", "GET /entidades/entidad-token + PDI", "200", "200 entidad PDI id 66", "200", "—", "PASS"),
                ("Endpoint inventado nodo", "GET /ddv3/endpoint-inexistente-qa", "404 ideal", "—", "500 vacío", "01KWJ8GK…", "OK + H-01"),
            ],
        ),
        (
            "3.2 — Test error 401: sin token JWT en Authorization",
            "EV-12",
            [
                ("Sin header Authorization", "GET /comunicaciones/89757", "401", "401 No autorizado", "500 vacío", "01KWHQW4…", "OK + H-01"),
            ],
        ),
        (
            "3.3 — Test error 401: token JWT mal formado o truncado",
            "EV-13, EV-14",
            [
                ("Token inválido", "Bearer token-invalido-123", "401", "401 No autorizado", "500 vacío", "01KWHRV58…", "OK + H-01"),
                ("Token truncado (~40 chars)", "Bearer fragmento JWT", "401", "401 No autorizado", "500 vacío", "—", "OK + H-01"),
            ],
        ),
        (
            "3.4 — Test error 401: token válido de otro sistema o ambiente",
            "EV-15",
            [
                ("Token ARMADA (otro sistema)", "GET /comunicaciones/89757 com. PDI", "401 ideal", "200 OK detalle completo", "200 OK", "01KWJB7N…", "H-03 — no reproduce 401"),
            ],
        ),
        (
            "3.5 — Test error 403: consulta detalle com. otra entidad",
            "EV-16",
            [
                ("ARMADA vs com. PDI 89757", "GET detalle", "403 ideal", "200 OK", "200 OK", "—", "H-03 — no reproduce 403"),
            ],
        ),
        (
            "3.6 — Test error 403: acuse recibo com. otra entidad",
            "EV-17",
            [
                ("Setup", "POST test-recepcion PDI → com. 89758, tarea 618993, folio NODO-005B", "200", "200", "200", "—", "PASS"),
                ("Acuse ARMADA", "PUT …/89758/recepcion/618993/acuse-recibo", "403", "403 Acceso denegado", "500 vacío", "01KWJDDDR…", "PASS directo + H-01 nodo"),
            ],
        ),
        (
            "3.7 — Test error 400: creación con tipos de datos incorrectos",
            "EV-18",
            [
                ("test-recepcion tipos mal", "folio=12345 (number), isReservado=\"texto\"", "400", "400 Petición no válida", "400", "—", "PASS — nodo propaga 400"),
            ],
        ),
        (
            "3.8 — Test error 400: búsqueda con filtros incorrectos",
            "EV-19",
            [
                ("buscar-entrantes", "?comunicacionId=abc", "400", "400 — abc no es Long", "500 vacío", "01KWJDRY5…", "OK + H-01"),
            ],
        ),
    ]

    for title, ev, casos in subtareas:
        heading(doc, title, 2)
        p = doc.add_paragraph()
        p.add_run("Evidencias: ").bold = True
        p.add_run(ev)

        t = doc.add_table(rows=1, cols=7)
        t.style = "Table Grid"
        headers = ["Caso", "Endpoint / acción", "Esp.", "Directo", "Nodo", "Id_traza", "Estado"]
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
            shade(t.rows[0].cells[i], "1F4E79")
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
        for caso in casos:
            row = t.add_row().cells
            for i, val in enumerate(caso):
                row[i].text = str(val)
                for run in row[i].paragraphs[0].runs:
                    run.font.size = Pt(8)
        doc.add_paragraph()

    doc.add_page_break()

    # 4 Matriz resumen
    heading(doc, "4. Matriz resumen subtareas", 1)
    resumen = [
        ("1", "404 inexistente / ruta incorrecta", "OK", "H-01 en errores"),
        ("2", "401 sin token", "OK", "H-01 nodo 500"),
        ("3", "401 inválido / truncado", "OK", "H-01 nodo 500"),
        ("4", "401 otro sistema (ARMADA)", "Ejecutado", "H-03 — 200, no 401"),
        ("5", "403 detalle (ARMADA)", "Ejecutado", "H-03 — 200, no 403"),
        ("6", "403 acuse (ARMADA)", "OK", "403 directo; H-01 nodo 500"),
        ("7", "400 creación", "OK", "400 directo y nodo"),
        ("8", "400 búsqueda", "OK", "400 directo; H-01 nodo 500"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["#", "Subtarea", "Ejecución", "Observación"]):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i], "1F4E79")
    for row in resumen:
        r = t.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    doc.add_page_break()

    # 5 Conclusiones
    heading(doc, "5. Conclusiones y recomendaciones", 1)

    bullets = [
        "La integración funcional Nodo PISEE ↔ DocDigital TEST está validada (OAuth, recepción, pendientes, acuse, consulta 200).",
        "Las 8 subtareas de manejo de errores fueron ejecutadas con evidencia comparativa directo vs nodo.",
        "Se debe escalar H-01 a Interoperabilidad/PISEE: el nodo debe propagar status y body de error (401/403/404/400) o documentar oficialmente el comportamiento 500.",
        "Confirmar con negocio H-03: ¿es esperado que ARMADA pueda leer GET de com. PDI en TEST? El acuse sí está protegido (403).",
        "Actualizar guías: reemplazar /entidades/token por /entidades/entidad-token.",
        "Tarjeta Asana: adjuntar Word evidencias + este informe. Comunicaciones clave: 89757, 89758 | tareas 618990, 618993.",
    ]
    for b in bullets:
        bullet(doc, b)

    doc.add_page_break()

    # 6 Comentario Asana
    heading(doc, "6. Texto sugerido — Comentario Asana (copiar y pegar)", 1)

    asana = """### Hallazgos (prioridad)

**H-01 — CRÍTICO:** Nodo no propaga errores HTTP DocDigital (401/403/404/400) → responde 500 body vacío. Escalar con Id_traza.

**H-03:** GET ARMADA vs com. PDI → 200 (no 401/403 detalle). PUT acuse ARMADA vs com. PDI → 403 directo OK; nodo 500.

### Flujo integración PASS
OAuth PDI → test-recepcion 89757 → pendientes 618990 → acuse 200 → GET 89757 200/200 + correo DocDigital.

### Subtareas error (8/8 ejecutadas)
| # | Resultado directo | Nodo | Nota |
|---|-------------------|------|------|
| 1 404 | 404 | 500 | H-01 |
| 2 401 sin token | 401 | 500 | H-01 |
| 3 401 inválido | 401 | 500 | H-01 |
| 4 ARMADA GET | 200 | 200 | H-03 |
| 5 ARMADA detalle | 200 | 200 | H-03 |
| 6 ARMADA acuse | 403 | 500 | OK + H-01 |
| 7 400 creación | 400 | 400 | OK |
| 8 400 búsqueda | 400 | 500 | H-01 |

**Adjuntos:** Documento_Evidencias_Pruebas NODO PISSE.docx + Informe QA 02-07-2026
**Ejecutor:** JFC | **Ambiente:** TEST"""

    p = doc.add_paragraph(asana)
    for run in p.runs:
        run.font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
