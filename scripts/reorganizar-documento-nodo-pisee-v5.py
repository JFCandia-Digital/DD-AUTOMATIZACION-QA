#!/usr/bin/env python3
"""Reorganiza Documento v5 del usuario en Word limpio y ordenado."""

import re
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/ubuntu/.cursor/projects/workspace/uploads/Documento_Evidencias_Pruebas_NODO_PISSE__5__fa5c.docx"
OUT = "/workspace/docs/api-v3.5/Documento_Evidencias_Pruebas_NODO_PISEE_DocDigital-ORDENADO.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def extract_images_from_para(p):
    imgs = []
    for run in p.runs:
        for blip in run._element.xpath(".//a:blip"):
            embed = blip.get(qn("r:embed"))
            if embed:
                imgs.append(embed)
    return imgs


def get_image_bytes(doc, rId):
    part = doc.part.related_parts[rId]
    return part.blob, part.content_type


def add_image_from_rid(doc, src_doc, rId, width_cm=16):
    blob, _content_type = get_image_bytes(src_doc, rId)
    stream = BytesIO(blob)
    doc.add_picture(stream, width=Cm(width_cm))


def add_heading(doc, text, level=2):
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    return p


def add_kv(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for r in table.rows[i].cells[0].paragraphs[0].runs:
            r.bold = True
    return table


def add_compare_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        set_cell_shading(c, "1F4E79")
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9)
    return table


def collect_sections(src):
    current_ev = "INIT"
    sections = {}
    order = []

    for p in src.paragraphs:
        t = p.text.strip()
        m = re.match(r"(EV-\d+)", t)
        if m:
            current_ev = m.group(1)
            if current_ev not in sections:
                sections[current_ev] = []
                order.append(current_ev)
        imgs = extract_images_from_para(p)
        if imgs:
            if current_ev not in sections:
                sections[current_ev] = []
                if current_ev not in order:
                    order.append(current_ev)
            sections[current_ev].extend(imgs)

    init_imgs = []
    for p in src.paragraphs[:3]:
        init_imgs.extend(extract_images_from_para(p))

    ev16_idxs = [114, 127, 129]
    ev17_idxs = [132, 139, 144, 153, 156]
    ev16_imgs = []
    ev17_imgs = []
    for idx in ev16_idxs:
        if idx < len(src.paragraphs):
            ev16_imgs.extend(extract_images_from_para(src.paragraphs[idx]))
    for idx in ev17_idxs:
        if idx < len(src.paragraphs):
            ev17_imgs.extend(extract_images_from_para(src.paragraphs[idx]))

    return sections, init_imgs, ev16_imgs, ev17_imgs


def section(doc, src, ev_title, desc_lines, result_table=None, image_rids=None):
    add_heading(doc, ev_title, 2)
    for line in desc_lines:
        if line == "":
            doc.add_paragraph()
        else:
            add_body(doc, line)
    if result_table:
        doc.add_paragraph()
        add_compare_table(doc, result_table[0], result_table[1])
    if image_rids:
        doc.add_paragraph()
        add_body(doc, "Evidencias (capturas):", bold=True)
        for rid in image_rids:
            try:
                add_image_from_rid(doc, src, rid)
                doc.add_paragraph()
            except Exception as exc:
                add_body(doc, f"[Imagen no insertada: {exc}]")
    doc.add_paragraph()


def main():
    src = Document(SRC)
    sections, init_imgs, ev16_imgs, ev17_imgs = collect_sections(src)

    def imgs_for(ev_key):
        if ev_key == "EV-16" and ev16_imgs:
            return ev16_imgs
        if ev_key == "EV-17" and ev17_imgs:
            return ev17_imgs
        return sections.get(ev_key, [])

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold in [
        ("DOCUMENTO DE EVIDENCIAS DE PRUEBAS\n", 16, True),
        ("Integración Nodo PISEE Consumidor ↔ DocDigital API v3.5\n", 14, True),
        ("Pruebas manuales vía NodoV2 + comparativo API directa TEST\n\n", 12, False),
        ("Ejecutor: Juan Francisco Candia (JFC) | Fecha: 02-07-2026\n", 11, False),
        ("Ambiente: TEST | Nodo PE-TEST-00001 :8085/ddv3\n", 11, False),
    ]:
        r = title.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    doc.add_page_break()

    add_heading(doc, "1. Configuración y alcance", 1)
    add_kv(
        doc,
        [
            ("Nodo PISEE", "v2.4.0 | Organismo PE-TEST-00001 | Puerto 8085 | Ruta /ddv3/*"),
            ("API DocDigital", "https://middleware.docv3.test.digital.gob.cl/api/v3"),
            ("Entidades", "PDI (66) | ARMADA (210)"),
            ("Método", "cURL manual — comparativo Directo vs Nodo"),
        ],
    )

    doc.add_paragraph()
    add_heading(doc, "2. Hallazgos principales", 1)
    add_body(
        doc,
        "H-01 (CRÍTICO): Errores 401/403/404/400 de la API directa → Nodo responde HTTP 500 body vacío (requiere Id_traza).",
    )
    add_body(
        doc,
        "H-03: GET cross-entidad en TEST retorna 200 (no 401/403 detalle). PUT acuse cross-entidad retorna 403 en API directa.",
    )

    doc.add_page_break()
    add_heading(doc, "3. Evidencias por caso de prueba", 1)

    if init_imgs:
        add_heading(doc, "EV-00 — Nodo levantado", 2)
        add_body(doc, "NodoV2 en ejecución — endpoint TEST visible.")
        add_image_from_rid(doc, src, init_imgs[0])
        doc.add_paragraph()

    section(
        doc,
        src,
        "EV-01 — Catálogo Nodo DocDigital v3",
        ["Pantalla CMD con endpoint TEST del catálogo Nodo."],
        image_rids=imgs_for("EV-01"),
    )
    section(
        doc,
        src,
        "EV-02 — OAuth token PDI vía Nodo (200)",
        ["POST /ddv3/oauth/token con credenciales PDI.", "Resultado: HTTP 200 + access_token Bearer."],
        (["Canal", "HTTP", "Resultado"], [["Nodo", "200", "PASS"]]),
        imgs_for("EV-02"),
    )
    section(
        doc,
        src,
        "EV-03 — test-recepcion vía Nodo (200)",
        ["POST test-recepcion — creación com. PDI."],
        (["Canal", "HTTP", "comunicacionId"], [["Nodo", "200", "89757"]]),
        imgs_for("EV-03"),
    )
    section(
        doc,
        src,
        "EV-04 — pendientes-recepcion vía Nodo (200)",
        ["GET pendientes-recepcion — obtiene tareaId para acuse."],
        (["Canal", "HTTP"], [["Nodo", "200"]]),
        imgs_for("EV-04"),
    )
    section(
        doc,
        src,
        "EV-05 — acuse-recibo vía Nodo (200)",
        ["PUT acuse-recibo — flujo recepción PDI."],
        (["Canal", "HTTP"], [["Nodo", "200"]]),
        imgs_for("EV-05"),
    )
    section(
        doc,
        src,
        "EV-06 — Correo DocDigital recepción",
        ["Correo de notificación.", "com. 89757 | folio QA-REC-JFC-NODO-003"],
        image_rids=imgs_for("EV-06"),
    )
    section(
        doc,
        src,
        "EV-07 — GET detalle 200 (directo vs nodo)",
        ["GET /comunicaciones/89757 con token PDI.", "Comparativo directo vs nodo — ambos 200."],
        (["Canal", "HTTP"], [["Directo", "200"], ["Nodo", "200"]]),
        imgs_for("EV-07"),
    )
    section(
        doc,
        src,
        "EV-08 — 404 comunicación inexistente",
        ["GET /comunicaciones/999999."],
        (
            ["Canal", "HTTP", "Body"],
            [["Directo", "404", "Registro no encontrado"], ["Nodo", "500", "vacío + Id_traza"]],
        ),
        imgs_for("EV-08"),
    )
    section(
        doc,
        src,
        "EV-09 — 404 ruta incorrecta /entidades/token",
        ["Ruta errónea vs API."],
        (["Canal", "HTTP"], [["Directo", "404"], ["Nodo", "500"]]),
        imgs_for("EV-09"),
    )
    section(
        doc,
        src,
        "EV-10 — Ruta correcta /entidades/entidad-token",
        ["Ruta correcta del catálogo."],
        (["Canal", "HTTP"], [["Directo", "200"], ["Nodo", "200"]]),
        imgs_for("EV-10"),
    )
    section(
        doc,
        src,
        "EV-11 — 404 endpoint inventado vía Nodo",
        ["GET /endpoint-inexistente-qa"],
        (["Canal", "HTTP"], [["Nodo", "500"]]),
        imgs_for("EV-11"),
    )
    section(
        doc,
        src,
        "EV-12 — 401 sin token",
        ["GET sin header Authorization."],
        (["Canal", "HTTP"], [["Directo", "401"], ["Nodo", "500"]]),
        imgs_for("EV-12"),
    )
    section(
        doc,
        src,
        "EV-13 — 401 token inválido",
        ["Authorization: Bearer token-invalido-123"],
        (["Canal", "HTTP"], [["Directo", "401"], ["Nodo", "500"]]),
        imgs_for("EV-13"),
    )
    section(
        doc,
        src,
        "EV-14 — 401 token truncado",
        ["Token JWT cortado a ~40 caracteres."],
        (["Canal", "HTTP"], [["Directo", "401"], ["Nodo", "500"]]),
        imgs_for("EV-14"),
    )
    section(
        doc,
        src,
        "EV-15 — Subtarea Asana #4 — Token otro sistema (ARMADA)",
        [
            "Token ARMADA (entidad 210) sobre com. PDI 89757.",
            "HTTP Directo: 200 OK (esperado 401 — NO reproducido)",
            "HTTP Nodo: 200 OK",
            "Id_traza nodo: 01KWJB7NAX7G63020GFMS0N35X",
            "Estado: EJECUTADO — HALLAZGO H-03",
        ],
        (["Canal", "HTTP", "Estado"], [["Directo", "200", "H-03"], ["Nodo", "200", "H-03"]]),
        imgs_for("EV-15"),
    )
    section(
        doc,
        src,
        "EV-16 — Subtarea Asana #5 — 403 detalle otra entidad",
        [
            "Paso 1 — Crear com. ARMADA (test-recepcion)",
            "  com. 89760 | folio QA-ARMADA-NODO-002 | HTTP 200",
            "  Id_traza: PE-TEST-00001.01KWJHPDTGCK1JS2JJ95K4SFMX",
            "",
            "Paso 2 — GET detalle con token PDI sobre com. ARMADA 89760",
            "  Token: PDI (entidad 66) | Esperado ideal: 403 | Obtenido: 200",
            "  HTTP Directo: 200 OK — detalle completo visible",
            "  HTTP Nodo: 200 OK | Id_traza: PE-TEST-00001.01KWJJ63T3RRK5EFVQRMH2TVHY",
            "",
            "Estado: EJECUTADO — HALLAZGO H-03 (403 detalle no reproducido en TEST)",
        ],
        (["Canal", "HTTP", "Estado"], [["Directo", "200", "H-03"], ["Nodo", "200", "H-03"]]),
        ev16_imgs,
    )
    section(
        doc,
        src,
        "EV-17 — Subtarea Asana #6 — 403 acuse otra entidad",
        [
            "com. 89758 | tarea 618993 | folio QA-REC-JFC-NODO-005B",
            "Paso 1 — Creación PDI: HTTP 200",
            "Paso 2 — Pendientes PDI: tareaId 618993 — HTTP 200",
            "Paso 3 — PUT acuse con token ARMADA:",
            "  Directo: HTTP 403 — Acceso denegado (PASS)",
            "  Nodo: HTTP 500 vacío — Id_traza: 01KWJDDDRJ2VGH75MTHPX675C5",
            "Hallazgo H-01: nodo no propaga 403",
            "Estado: PASS directo + HALLAZGO nodo (H-01)",
        ],
        (["Canal", "HTTP", "Estado"], [["Directo", "403", "PASS"], ["Nodo", "500", "H-01"]]),
        ev17_imgs,
    )
    section(
        doc,
        src,
        "EV-18 — Subtarea Asana #7 — 400 creación tipos incorrectos",
        [
            "POST test-recepcion con tipos de dato inválidos.",
            "HTTP Directo: 400 | HTTP Nodo: 400",
            "Estado: PASS",
        ],
        (["Canal", "HTTP"], [["Directo", "400"], ["Nodo", "400"]]),
        imgs_for("EV-18"),
    )
    section(
        doc,
        src,
        "EV-19 — Subtarea Asana #8 — 400 búsqueda filtro inválido",
        [
            "GET buscar-entrantes?comunicacionId=abc",
            "HTTP Directo: 400 — Petición no válida",
            "HTTP Nodo: 500 — Id_traza: 01KWJDRY5XK0AESZE00QX2WPHV",
            "Estado: PASS directo + HALLAZGO H-01 nodo",
        ],
        (["Canal", "HTTP"], [["Directo", "400"], ["Nodo", "500"]]),
        imgs_for("EV-19"),
    )

    doc.add_page_break()
    add_heading(doc, "4. Resumen — Matriz subtareas Asana (EV-20)", 1)
    add_body(doc, "Fecha: 02-07-2026 | Ejecutor: JFC | Ambiente TEST")
    doc.add_paragraph()
    add_compare_table(
        doc,
        ["#", "Subtarea Asana", "Directo", "Nodo", "Estado"],
        [
            ["1", "404 no encontrado", "404", "500", "OK + H-01"],
            ["2", "401 sin token", "401", "500", "OK + H-01"],
            ["3", "401 inválido/truncado", "401", "500", "OK + H-01"],
            ["4", "401 otro sistema", "200", "200", "H-03"],
            ["5", "403 detalle", "200", "200", "H-03"],
            ["6", "403 acuse", "403", "500", "OK + H-01"],
            ["7", "400 creación", "400", "400", "OK"],
            ["8", "400 búsqueda", "400", "500", "OK + H-01"],
        ],
    )

    doc.add_paragraph()
    add_heading(doc, "5. Comentario sugerido — Asana", 1)
    asana_text = (
        "Pruebas integración Nodo PISEE v2.4.0 ↔ DocDigital API v3.5 (TEST) completadas.\n\n"
        "Flujo E2E: PASS (OAuth, test-recepcion, pendientes, acuse, GET detalle).\n\n"
        "8 subtareas error ejecutadas:\n"
        "• #1-3, #8: API directa OK; nodo devuelve 500 vacío (H-01).\n"
        "• #4, #5: token cross-entidad GET → 200/200 (H-03).\n"
        "• #6: acuse cross-entidad → 403 directo OK; nodo 500 (H-01).\n"
        "• #7: 400 creación → 400/400 (OK).\n\n"
        "Comunicaciones: 89757, 89758, 89760 | tarea 618993.\n"
        "Evidencias en documento adjunto."
    )
    p = doc.add_paragraph(asana_text)
    for r in p.runs:
        r.font.size = Pt(9)

    doc.save(OUT)
    print(f"Generado: {OUT}")
    print(f"Imágenes EV-16: {len(ev16_imgs)}, EV-17: {len(ev17_imgs)}")


if __name__ == "__main__":
    main()
