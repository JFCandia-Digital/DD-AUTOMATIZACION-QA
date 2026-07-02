#!/usr/bin/env python3
"""Genera Documento_Evidencias_Pruebas_Nodo_PISEE_DocDigital.docx"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/api-v3.5/Documento_Evidencias_Pruebas_Nodo_PISEE_DocDigital.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    return table


def add_data_table(doc, headers, rows, header_color="1F4E79", font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(font_size)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    return table


def shade_status_column(table, col_idx, pass_values=("PASS", "200", "OK", "Sí")):
    for row in table.rows[1:]:
        cell = row.cells[col_idx]
        text = cell.text.strip().upper()
        if any(v.upper() in text for v in pass_values):
            set_cell_shading(cell, "C6EFCE")
        elif "FALLO" in text or "500" in text or "HALLAZGO" in text:
            set_cell_shading(cell, "FFC7CE")
        elif "PEND" in text or "OBS" in text:
            set_cell_shading(cell, "FFEB9C")


def main():
    doc = Document()

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold, italic in [
        ("DOCUMENTO DE EVIDENCIAS DE PRUEBAS\n", 16, True, False),
        ("Integración Nodo PISEE Consumidor ↔ DocDigital API v3.5\n", 14, False, False),
        ("Pruebas manuales vía NodoV2 + comparativo API directa TEST\n\n", 12, False, False),
        ("Tarjeta Asana — Integración Nodo PISEE consumidor / DocDigital API v3.5", 11, False, True),
    ]:
        run = title.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    doc.add_paragraph()
    add_kv_table(doc, [
        ("Proyecto", "DD-AUTOMATIZACION-QA + Nodo PISEE local (NODOCONSTEST)"),
        ("Módulo probado", "Nodo consumidor DocDigitalv3 — proxy /ddv3/* → middleware TEST"),
        ("Ambiente efectivo", "TEST — https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Versión API", "Middleware DocDigital v3.5"),
        ("Nodo", "NodoV2.exe v2.4.0 — organismo PE-TEST-00001 — puerto 8085"),
        ("Período de pruebas", "30-06-2026 al 02-07-2026"),
        ("Ejecutado por", "Juan Francisco Candia (JFC) — QA Automatizador"),
        ("Repositorio QA", "JFCandia-Digital/DD-AUTOMATIZACION-QA"),
        ("Repositorio nodo", "JFCandia-Digital/NODOCONSTEST"),
        ("Resultado global", "Flujo principal PASS | Hallazgo: errores 401/404 no propagados por nodo"),
    ])

    doc.add_page_break()

    # 1. Hallazgos primero (formato acordado Asana)
    add_heading(doc, "1. Hallazgos y observaciones (prioridad)", 1)

    hallazgos = [
        ("H-01", "Alta", "El nodo PISEE no propaga códigos HTTP de error de DocDigital API. "
         "Ante 401 (sin token, token inválido) y 404 (recurso inexistente) en API directa, "
         "el nodo responde 500 Internal Server Error con body vacío."),
        ("H-02", "Media", "Ruta incorrecta documentada en guías antiguas: GET /entidades/token "
         "retorna 404 directo y 500 vía nodo. La ruta correcta según Swagger es "
         "GET /entidades/entidad-token (200 directo y vía nodo con token PDI)."),
        ("H-03", "Baja", "Caso 403 no reproducido: token ARMADA (entidad distinta) sobre "
         "comunicación 89754 retorna 200 en ambos canales, no 403 Forbidden."),
        ("H-04", "Info", "Tarjeta Asana indica ambiente DEMO; catálogo del nodo apunta a TEST. "
         "Equipo confirmó que TEST alcanza para las pruebas de integración."),
        ("H-05", "Info", "GET /entidades/token (ruta vieja) vía nodo → 500 en sesión 30-06; "
         "no bloquea el flujo si test-recepcion y pendientes responden 200."),
    ]
    t_h = add_data_table(doc, ["ID", "Severidad", "Descripción"], hallazgos)
    shade_status_column(t_h, 1, ("ALTA", "MEDIA", "BAJA", "INFO"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Impacto H-01: ").bold = True
    p.add_run(
        "Los integradores que consuman DocDigital solo vía nodo no reciben el código ni el mensaje "
        "de error real (401 No autorizado, 404 Registro no encontrado). Deben usar Id_traza PISEE "
        "para escalar a Interoperabilidad. Relacionado con solicitud Patricia (24-04) sobre "
        "latencia y manejo de excepciones."
    )

    doc.add_page_break()

    # 2. Objetivo
    add_heading(doc, "2. Objetivo", 1)
    doc.add_paragraph(
        "Documentar las pruebas de integración entre el Nodo PISEE consumidor (servicio catálogo "
        "DocDigitalv3) y la API Middleware DocDigital v3.5 en ambiente TEST, comparando respuestas "
        "vía nodo local (localhost:8085/ddv3/*) frente a llamadas directas al middleware, "
        "incluyendo flujo funcional completo de recepción y casos de borde de seguridad y errores."
    )

    # 3. Alcance
    add_heading(doc, "3. Alcance", 1)
    for item in [
        "Verificación de catálogo y conectividad (NodoV2 servicios DocDigitalv3).",
        "Autenticación OAuth2 vía nodo (POST /ddv3/oauth/token).",
        "Flujo funcional: test-recepcion → pendientes-recepcion → acuse-recibo.",
        "Consultas GET comunicación y trazabilidad (happy path 200).",
        "Casos de seguridad: token ausente, token inválido (Bearer falso).",
        "Casos de error: comunicación inexistente (404).",
        "Intento de autorización cruzada (token ARMADA vs comunicación PDI).",
        "Corrección de ruta entidades (entidad-token vs token).",
        "Comparativo directo vs nodo en todos los casos anteriores.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 4. Configuración
    add_heading(doc, "4. Configuración y metodología de prueba", 1)

    add_heading(doc, "4.1 Configuración del nodo PISEE", 2)
    add_kv_table(doc, [
        ("Carpeta nodo", r"C:\Users\Juan Candia\Documents\3 - Api_NodoConsTEST_DocDigital_dev\NodoConsTEST"),
        ("Ejecutable", "NodoV2.exe v2.4.0"),
        ("Organismo", "PE-TEST-00001"),
        ("Puerto interno", "8085"),
        ("Ruta local", "/ddv3/*"),
        ("Servicio catálogo", "DocDigitalv3"),
        ("Endpoint catálogo", "https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Puerto externo PISEE", "8490"),
    ])

    add_heading(doc, "4.2 Cómo se probó (procedimiento)", 2)
    doc.add_paragraph(
        "Todas las pruebas se ejecutaron desde dos ventanas CMD en Windows 10, siguiendo el manual "
        "Nodo v2 PISEE. El nodo se mantuvo levantado en CMD 1 (NodoV2.exe start) y las peticiones "
        "curl se lanzaron desde CMD 2 en la misma carpeta del nodo."
    )
    pasos = [
        ("Paso 1", "Verificar catálogo: NodoV2.exe servicios DocDigitalv3 — confirmar Endpoint TEST."),
        ("Paso 2", "Obtener token OAuth vía nodo: POST http://localhost:8085/ddv3/oauth/token con client_id:client_secret."),
        ("Paso 3", "Ejecutar flujo funcional con token PDI (entidad 66) en endpoints /ddv3/..."),
        ("Paso 4", "Repetir casos seleccionados contra API directa: https://middleware.docv3.test.digital.gob.cl/api/v3/..."),
        ("Paso 5", "Comparar HTTP status, body JSON, cabecera Id_traza (solo nodo) y correo DocDigital."),
        ("Paso 6", "Registrar evidencias: capturas curl, Id_traza, comunicacionId, tareaId, folio."),
    ]
    add_data_table(doc, ["#", "Acción"], pasos)

    doc.add_paragraph()
    add_heading(doc, "4.3 Tokens utilizados por escenario", 2)
    doc.add_paragraph(
        "Es importante distinguir qué token corresponde a cada caso. Usar token PDI en pruebas de "
        "error de recurso (404) es correcto; usar token inválido solo en pruebas de autenticación (401)."
    )
    tokens = [
        ("PDI (válido)", "OAuth con credenciales PDI", "Flujos 200, 404, E2E recepción", "Entidad 66 — Policía de Investigaciones"),
        ("Ausente (nulo)", "Sin cabecera Authorization", "Prueba 401 — token no enviado", "Caso TC-INT-008 sub-caso nulo"),
        ("token-invalido-123", "Bearer token-invalido-123", "Prueba 401 — token inválido (automation Cucumber)", "Mismo valor que apiClient.ts"),
        ("token_invalido", "Bearer token_invalido", "Prueba 401 — variante manual (opcional)", "Mismo comportamiento esperado que guión"),
        ("ARMADA (válido)", "OAuth credenciales ARMADA", "Intento 403 — otra entidad", "Token válido, no inválido"),
    ]
    add_data_table(doc, ["Tipo", "Header / origen", "Cuándo usar", "Notas"], tokens)

    doc.add_page_break()

    # 5. Flujo principal PASS
    add_heading(doc, "5. Flujo principal — Resultado PASS", 1)

    add_heading(doc, "5.1 Sesión 30-06-2026 (primera integración)", 2)
    add_data_table(doc,
        ["Paso", "Endpoint vía nodo", "Método", "Token", "HTTP", "Resultado"],
        [
            ("1", "/ddv3/oauth/token", "POST", "Basic OAuth", "200", "Token PDI obtenido"),
            ("2", "/ddv3/entidades/token", "GET", "PDI", "500", "Body vacío — ruta incorrecta (ver H-02)"),
            ("3", "/ddv3/pruebas-integracion/comunicaciones/test-recepcion", "POST", "PDI", "200", "com. 89753 creada"),
            ("4", "/ddv3/comunicaciones/pendientes-recepcion", "GET", "PDI", "200", "tarea 618966 listada"),
        ],
    )
    doc.add_paragraph()
    add_kv_table(doc, [
        ("comunicacionId", "89753"),
        ("tareaId", "618966"),
        ("folio", "QA-REC-JFC-NODO-001"),
        ("materia", "QA Test Recepcion JFC nodo"),
        ("fechaDespacho", "30-06-2026 14:39:00"),
        ("despachadora", "Entidad TEST Agosto KE"),
        ("Correo DocDigital", "Recibido — com. 89753"),
        ("Id_traza (ej.)", "PE-TEST-00001.01KWCX7JTTD3DY2NAGDJJ7M80H (test-recepcion)"),
    ])

    add_heading(doc, "5.2 Sesión 01-07-2026 (E2E completo)", 2)
    add_data_table(doc,
        ["Paso", "Endpoint", "Método", "HTTP nodo", "Resultado"],
        [
            ("1", "/ddv3/oauth/token", "POST", "200", "Token PDI"),
            ("2", "/ddv3/.../test-recepcion", "POST", "200", "com. 89754 — folio QA-REC-JFC-NODO-002"),
            ("3", "/ddv3/comunicaciones/pendientes-recepcion", "GET", "200", "tarea 618969"),
            ("4", "/ddv3/comunicaciones/89754/recepcion/618969/acuse-recibo", "PUT", "200", "Acuse isRechazada:false"),
            ("5", "Correo DocDigital", "—", "—", "Recibido — com. 89754"),
        ],
    )
    doc.add_paragraph()
    add_kv_table(doc, [
        ("comunicacionId", "89754"),
        ("tareaId", "618969"),
        ("folio", "QA-REC-JFC-NODO-002"),
        ("Payload acuse", '{"isRechazada":false,"comentarios":"QA acuse vía nodo PISEE"}'),
    ])

    doc.add_paragraph()
    add_heading(doc, "5.3 Payload test-recepcion utilizado", 2)
    doc.add_paragraph(
        '{\n'
        '  "materia": "QA Test Recepcion JFC nodo",\n'
        '  "folio": "QA-REC-JFC-NODO-00X",\n'
        '  "isReservado": false,\n'
        '  "incorporaAnexos": false,\n'
        '  "asociarProcedimientoAdministrativo": false\n'
        '}'
    )

    doc.add_page_break()

    # 6. Tabla comparativa completa
    add_heading(doc, "6. Tabla comparativa — API directa TEST vs Nodo PISEE", 1)
    doc.add_paragraph(
        "Cada fila resume una ejecución manual con curl. Columna «Body resumido» indica el mensaje "
        "clave de la respuesta. Cuando el nodo devuelve 500, el body viene vacío y solo queda "
        "la cabecera Id_traza para trazabilidad PISEE."
    )

    comparativo = [
        ("C-01", "GET /comunicaciones/89754", "PDI válido", "200", "200 OK — detalle com.", "200 OK", "—", "PASS"),
        ("C-02", "GET /comunicaciones/89754/trazabilidad", "PDI válido", "200", "200 OK — trazabilidad", "200 OK", "—", "PASS"),
        ("C-03", "GET /comunicaciones/89754", "Sin Authorization", "401", "No autorizado.", "500 vacío", "01KWHQW4C9YWNKA7DS9NY813SS", "HALLAZGO H-01"),
        ("C-04", "GET /comunicaciones/89754", "Bearer token-invalido-123", "401", "No autorizado.", "500 vacío", "01KWHRV58B30PSF8RS5MBZNPFH", "HALLAZGO H-01"),
        ("C-05", "GET /comunicaciones/89754", "Bearer token_invalido", "401 (esp.)", "No autorizado. (esp.)", "500 vacío (esp.)", "Pendiente captura", "Pendiente / mismo patrón"),
        ("C-06", "GET /comunicaciones/999999", "PDI válido", "404", "Registro no encontrado.", "500 vacío", "01KWHR3T9SGSBZXNRWJ6K7GT35", "HALLAZGO H-01"),
        ("C-07", "GET /comunicaciones/999999/trazabilidad", "PDI válido", "404", "Registro no encontrado.", "500 vacío", "01KWHR7H0TRTQS48PAMW7XKQPA", "HALLAZGO H-01"),
        ("C-08", "GET /comunicaciones/89754", "ARMADA válido", "200", "200 (no 403)", "200 (no 403)", "—", "OBS H-03"),
        ("C-09", "GET /entidades/token", "PDI válido", "404", "Ruta no existe", "500 vacío", "—", "FALLO ruta"),
        ("C-10", "GET /entidades/entidad-token", "PDI válido", "200", "Entidad PDI id 66", "200 OK", "—", "PASS (ruta correcta)"),
        ("C-11", "POST /oauth/token", "Basic OAuth PDI", "200", "access_token", "200 OK", "—", "PASS"),
        ("C-12", "POST .../test-recepcion", "PDI válido", "200", "result.id", "200 — com. 89754", "—", "PASS"),
        ("C-13", "GET /pendientes-recepcion", "PDI válido", "200", "lista tareas", "200 — tarea 618969", "—", "PASS"),
        ("C-14", "PUT .../acuse-recibo", "PDI válido", "200", "acuse completado", "200 OK", "—", "PASS"),
    ]
    t_comp = add_data_table(
        doc,
        ["ID", "Endpoint", "Token", "HTTP Directo", "Body directo", "HTTP Nodo", "Id_traza nodo", "Estado"],
        comparativo,
        font_size=7,
    )
    shade_status_column(t_comp, 7, ("PASS",))
    for row in t_comp.rows[1:]:
        estado = row.cells[7].text
        if "HALLAZGO" in estado:
            set_cell_shading(row.cells[7], "FFC7CE")
        elif "FALLO" in estado or "OBS" in estado:
            set_cell_shading(row.cells[7], "FFEB9C")
        elif "PEND" in estado:
            set_cell_shading(row.cells[7], "FFEB9C")

    doc.add_page_break()

    # 7. Comandos curl ejecutados
    add_heading(doc, "7. Comandos curl de referencia (evidencia)", 1)

    comandos = [
        ("OAuth vía nodo",
         'curl -i -X POST "http://localhost:8085/ddv3/oauth/token" -u "CLIENT_ID:CLIENT_SECRET"'),
        ("test-recepcion vía nodo",
         'curl -i -X POST "http://localhost:8085/ddv3/pruebas-integracion/comunicaciones/test-recepcion" '
         '-H "Authorization: Bearer TOKEN_PDI" -H "Content-Type: application/json" '
         '-d "{\\"materia\\":\\"QA Test Recepcion JFC nodo\\",\\"folio\\":\\"QA-REC-JFC-NODO-002\\",'
         '\\"isReservado\\":false,\\"incorporaAnexos\\":false,\\"asociarProcedimientoAdministrativo\\":false}"'),
        ("pendientes vía nodo",
         'curl -i -X GET "http://localhost:8085/ddv3/comunicaciones/pendientes-recepcion" '
         '-H "Authorization: Bearer TOKEN_PDI"'),
        ("acuse-recibo vía nodo",
         'curl -i -X PUT "http://localhost:8085/ddv3/comunicaciones/89754/recepcion/618969/acuse-recibo" '
         '-H "Authorization: Bearer TOKEN_PDI" -H "Content-Type: application/json" '
         '-d "{\\"isRechazada\\":false,\\"comentarios\\":\\"QA acuse vía nodo\\"}"'),
        ("401 sin token — directo",
         'curl -i -X GET "https://middleware.docv3.test.digital.gob.cl/api/v3/comunicaciones/89754"'),
        ("401 sin token — nodo",
         'curl -i -X GET "http://localhost:8085/ddv3/comunicaciones/89754"'),
        ("401 token inválido — directo",
         'curl -i -X GET "https://middleware.docv3.test.digital.gob.cl/api/v3/comunicaciones/89754" '
         '-H "Authorization: Bearer token-invalido-123"'),
        ("401 token inválido — nodo",
         'curl -i -X GET "http://localhost:8085/ddv3/comunicaciones/89754" '
         '-H "Authorization: Bearer token-invalido-123"'),
        ("404 inexistente — directo",
         'curl -i -X GET "https://middleware.docv3.test.digital.gob.cl/api/v3/comunicaciones/999999" '
         '-H "Authorization: Bearer TOKEN_PDI"'),
        ("404 inexistente — nodo",
         'curl -i -X GET "http://localhost:8085/ddv3/comunicaciones/999999" '
         '-H "Authorization: Bearer TOKEN_PDI"'),
        ("entidad-token correcto — directo",
         'curl -i -X GET "https://middleware.docv3.test.digital.gob.cl/api/v3/entidades/entidad-token" '
         '-H "Authorization: Bearer TOKEN_PDI"'),
        ("entidad-token correcto — nodo",
         'curl -i -X GET "http://localhost:8085/ddv3/entidades/entidad-token" '
         '-H "Authorization: Bearer TOKEN_PDI"'),
    ]
    for nombre, cmd in comandos:
        p = doc.add_paragraph()
        p.add_run(f"{nombre}:\n").bold = True
        run = doc.add_paragraph(cmd)
        for r in run.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(8)

    doc.add_page_break()

    # 8. Evidencias
    add_heading(doc, "8. Checklist de evidencias (adjuntar en Asana)", 1)
    evidencias = [
        ("E-01", "Catálogo nodo", "NodoV2 servicios DocDigitalv3 — Endpoint TEST", "Captura CMD"),
        ("E-02", "OAuth 200", "POST /ddv3/oauth/token — token PDI", "Captura curl (tapar secret)"),
        ("E-03", "test-recepcion 200", "com. 89753 y 89754", "Captura curl + JSON result.id"),
        ("E-04", "pendientes 200", "tareas 618966 y 618969", "Captura curl"),
        ("E-05", "acuse-recibo 200", "PUT acuse com. 89754 tarea 618969", "Captura curl"),
        ("E-06", "Correos DocDigital", "89753, 89754", "Captura Gmail"),
        ("E-07", "Comparativo 401 sin token", "Directo 401 vs Nodo 500 + Id_traza", "Captura lado a lado"),
        ("E-08", "Comparativo 401 token inválido", "token-invalido-123 ambos canales", "Captura lado a lado"),
        ("E-09", "Comparativo 404", "999999 directo 404 vs nodo 500", "Captura lado a lado"),
        ("E-10", "entidad-token 200", "Ruta correcta Swagger", "Captura curl"),
        ("E-11", "Este documento Word", "Documento_Evidencias_Pruebas_Nodo_PISEE_DocDigital.docx", "Adjunto Asana"),
        ("E-12", "README guía", "docs/api-v3.5/README-Pruebas-Nodo-PISEE.md", "Repositorio QA"),
    ]
    add_data_table(doc, ["ID", "Tipo", "Descripción", "Formato"], evidencias)

    doc.add_paragraph()
    for label in [
        "E-07 — 401 sin Authorization (directo vs nodo):",
        "E-08 — 401 token-invalido-123 (directo vs nodo):",
        "E-09 — 404 comunicación 999999 (directo vs nodo):",
        "E-06 — Correos DocDigital 89753 y 89754:",
    ]:
        doc.add_paragraph(label).runs[0].bold = True
        doc.add_paragraph("[ Insertar captura aquí ]").runs[0].italic = True
        doc.add_paragraph()

    doc.add_page_break()

    # 9. Plantilla Asana
    add_heading(doc, "9. Plantilla — Comentario para Asana (copiar y pegar)", 1)
    doc.add_paragraph(
        "Formato acordado: (1) Errores/hallazgos al inicio, (2) Comparativo directo vs nodo, "
        "(3) Flujo PASS, (4) Config y payloads, (5) Evidencias numeradas."
    )

    asana_text = """### Hallazgos / errores (prioridad)

1. **H-01 — Nodo no propaga errores HTTP de DocDigital API**
   - 401 (sin token, token inválido) → nodo responde **500 body vacío**
   - 404 (comunicación inexistente) → nodo responde **500 body vacío**
   - API directa TEST responde correctamente 401/404 con JSON de error
   - Id_traza ejemplo: 01KWHQW4… (401 sin token), 01KWHRV58… (token inválido), 01KWHR3T… (404)

2. **H-02 — Ruta entidades incorrecta en guías**
   - GET /entidades/token → 404 directo / 500 nodo
   - GET /entidades/entidad-token → **200** directo y nodo (PDI id 66)

3. **H-03 — 403 no reproducido** con token ARMADA sobre com. 89754 (ambos 200)

---

### Comparativo directo vs nodo (02-07-2026)

| Caso | Directo TEST | Nodo localhost:8085 |
|------|--------------|---------------------|
| GET /comunicaciones/89754 + PDI | 200 | 200 |
| GET /comunicaciones/89754/trazabilidad + PDI | 200 | 200 |
| GET sin Authorization | 401 No autorizado | 500 vacío |
| GET Bearer token-invalido-123 | 401 No autorizado | 500 vacío |
| GET /comunicaciones/999999 + PDI | 404 Registro no encontrado | 500 vacío |
| GET /comunicaciones/999999/trazabilidad + PDI | 404 | 500 vacío |
| GET /comunicaciones/89754 + ARMADA | 200 (no 403) | 200 |
| GET /entidades/entidad-token + PDI | 200 | 200 |

---

### Flujo principal PASS (token PDI)

| Fecha | Paso | Resultado |
|-------|------|-----------|
| 30-06 | OAuth → test-recepcion → pendientes | 200 — com. **89753**, tarea **618966**, folio QA-REC-JFC-NODO-001 |
| 01-07 | OAuth → test-recepcion → pendientes → acuse-recibo | 200 — com. **89754**, tarea **618969**, folio QA-REC-JFC-NODO-002 |
| — | Correo DocDigital | Recibido en ambos casos |

---

### Configuración

- Catálogo: `https://middleware.docv3.test.digital.gob.cl/api/v3`
- Nodo: NodoV2 v2.4.0 | PE-TEST-00001 | puerto 8085 | ruta `/ddv3/*`
- Token flujo OK: **PDI** (entidad 66)
- Token 401 inválido: `Bearer token-invalido-123` (no es PDI)
- Token 401 ausente: sin header Authorization

---

### Evidencias adjuntas

- E-01 a E-12 — ver Documento Word `Documento_Evidencias_Pruebas_Nodo_PISEE_DocDigital.docx`
- Repositorio guía: `docs/api-v3.5/README-Pruebas-Nodo-PISEE.md`

---

**Ejecutor:** Juan Francisco Candia (JFC) — QA Automatizador
**Ambiente:** TEST (tarjeta decía DEMO; equipo confirmó TEST alcanza)
**Estado:** Flujo integración PASS | Hallazgo manejo errores nodo — escalar H-01 a Interoperabilidad/PISEE"""

    p_asana = doc.add_paragraph(asana_text)
    for r in p_asana.runs:
        r.font.size = Pt(9)
        r.font.name = "Consolas"

    doc.add_page_break()

    # 10. Conclusiones
    add_heading(doc, "10. Conclusiones", 1)
    doc.add_paragraph(
        "La integración funcional entre el Nodo PISEE consumidor DocDigitalv3 y la API Middleware "
        "v3.5 en ambiente TEST fue validada exitosamente: OAuth, creación de comunicación de "
        "prueba, consulta de pendientes y acuse de recibo responden HTTP 200 tanto vía nodo "
        "como en comparación con la API directa en escenarios de happy path."
    )
    doc.add_paragraph(
        "Se identificó un hallazgo relevante (H-01): ante errores de autenticación (401) y "
        "recurso no encontrado (404) en la API directa, el nodo transforma la respuesta en "
        "500 Internal Server Error sin cuerpo, impidiendo al consumidor interpretar el error real. "
        "Este comportamiento debe reportarse a Interoperabilidad/PISEE para corrección o "
        "documentación oficial."
    )
    doc.add_paragraph(
        "Se recomienda actualizar la guía de integración reemplazando /entidades/token por "
        "/entidades/entidad-token, y mantener la distinción clara entre token PDI (flujos 200), "
        "token ausente/inválido (pruebas 401) y token de otra entidad (pruebas 403)."
    )

    add_heading(doc, "11. Anexos", 1)
    for item in [
        "Anexo A: README-Pruebas-Nodo-PISEE.md — docs/api-v3.5/",
        "Anexo B: Repositorio NODOCONSTEST — github.com/JFCandia-Digital/NODOCONSTEST",
        "Anexo C: Manual Nodo v2 PISEE (PDF entregado por Pablo — 16 abr)",
        "Anexo D: Capturas curl y correos (insertar en sección 8)",
        "Anexo E: Swagger API MW v3.5 — ruta entidad-token",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.add_run("\n\n_______________________________\n")
    firma.add_run("Juan Francisco Candia (JFC)\n").bold = True
    firma.add_run("QA Automatizador — DocDigital\n")
    firma.add_run("Fecha documento: 02-07-2026")

    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
