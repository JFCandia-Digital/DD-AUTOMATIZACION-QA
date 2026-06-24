#!/usr/bin/env python3
"""Genera Documento_Evidencias_Pruebas_TestRecepcion.docx — QA-5687"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/api-v3.5/Documento_Evidencias_Pruebas_TestRecepcion.docx"

CASOS = [
    ("TC-INT-001", "API / Happy path", "Envío comunicación de prueba con materia y folio (caso mínimo)",
     "POST test-recepcion | JSON mínimo | Validar 200 + result.id",
     "200 OK, comunicación creada", "200 OK — id 89722 (24-06-2026 12:20:22)", "PASS",
     "reports/index.html; correo DocDigital 89722"),
    ("TC-INT-002", "API / Funcional", "Envío con incorporaAnexos=true",
     "POST con incorporaAnexos: true", "200 OK con anexos de prueba",
     "200 OK — id 89723", "PASS", "reports/index.html; correo 89723"),
    ("TC-INT-003", "API / Funcional", "Envío con isReservado=true",
     "POST con isReservado: true", "200 OK, flag reservado",
     "200 OK — id 89724", "PASS", "reports/index.html; correo 89724"),
    ("TC-INT-004", "API / Funcional", "Envío con asociarProcedimientoAdministrativo=true",
     "POST con PA asociado", "200 OK, PA de prueba",
     "200 OK — id 89725", "PASS", "reports/index.html; correo 89725"),
    ("TC-INT-005", "API / Funcional", "Envío con todos los flags opcionales en true",
     "POST flags combinados", "200 OK",
     "200 OK — id 89726", "PASS", "reports/index.html; correo 89726"),
    ("TC-INT-006", "API / Validación", "Rechazo sin materia",
     "POST sin campo materia", "400 Bad Request",
     "400 — errorCode 40000 — materia es obligatorio", "PASS", "reports/index.html"),
    ("TC-INT-007", "API / Validación", "Rechazo sin folio",
     "POST sin campo folio", "400 Bad Request",
     "400 — errorCode 40000 — folio es obligatorio", "PASS", "reports/index.html"),
    ("TC-INT-008", "API / Seguridad", "Token inválido o ausente",
     "POST con token inválido/nulo", "401 Unauthorized",
     "401 — No autorizado.", "PASS", "reports/index.html"),
    ("TC-INT-009", "API / Seguridad", "Token expirado",
     "POST con token expirado", "401 Unauthorized",
     "401 — Sesión expirada.", "PASS", "reports/index.html"),
    ("TC-INT-010", "API / Contrato", "Usuario no elige documento principal",
     "Verificar endpoint solo JSON (sin multipart)", "Sin campo archivo en API",
     "Pendiente — validación manual Swagger", "Pendiente", "Swagger spec"),
    ("TC-INT-011", "API / Negocio", "Mismo PDF con isReservado true vs false",
     "2 POST + comparar archivo", "Mismo documento principal",
     "Pendiente — GET metadata/hash", "Pendiente", ""),
    ("TC-INT-012", "API / Negocio", "Anexos idénticos entre envíos",
     "2 POST con anexos + comparar", "Mismos anexos fijos",
     "Pendiente — GET anexos", "Pendiente", ""),
    ("TC-INT-013", "API / Negocio", "Despachadora Agosto KE",
     "GET trazabilidad/detalle", "Entidad despachadora Agosto KE",
     "Pendiente — GET comunicación", "Pendiente", ""),
    ("TC-INT-014", "API / Negocio", "PA por defecto Otro procedimiento",
     "GET detalle con PA", "PA Otro procedimiento",
     "Pendiente — GET detalle", "Pendiente", ""),
    ("TC-INT-015", "API / Negocio", "Documento principal fijo con FEA",
     "GET metadata documento", "prueba_dipres_test.pdf con FEA",
     "Pendiente — GET archivo", "Pendiente", ""),
    ("TC-INT-016", "API / Funcional", "Creación individual (no masiva)",
     "1 POST = 1 id", "Un id por request",
     "Validado en suite (89722–89726 distintos)", "PASS", "Suite Cucumber"),
    ("TC-INT-017", "API / E2E", "Flujo acuse de recibo",
     "POST → GET pendientes → PUT acuse", "Flujo recepción completo",
     "Pendiente ejecución E2E — npm run test_recepcion_e2e", "Pendiente",
     "postTestRecepcionE2E.feature @TestRecepcion_E2E_Acuse"),
    ("TC-INT-018", "API / E2E", "Flujo rechazo",
     "POST → GET pendientes → PUT rechazo", "Flujo rechazo completo",
     "Pendiente ejecución E2E — npm run test_recepcion_e2e", "Pendiente",
     "postTestRecepcionE2E.feature @TestRecepcion_E2E_Rechazo"),
]

BLOQUEO = [
    ("23-06-2026", "HTTP 401 / errorCode 40100 — [ENV CREDENCIALES ENTIDAD_TEST]",
     "Variables servidor QA no configuradas"),
    ("24-06-2026", "Susana solicita config a Pablo — API_CLIENT_ENTIDAD_TEST / API_SECRET_ENTIDAD_TEST",
     "Credenciales entidad test en INFRA QA"),
    ("24-06-2026", "POST test-recepcion → HTTP 200 — ids 89721 (probe), 89722–89726 (suite)",
     "Tarjeta destrabada — automatización ejecutada"),
]


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    return table


def main():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENTO DE EVIDENCIAS DE PRUEBAS\n")
    run.bold = True
    run.font.size = Pt(16)
    run2 = title.add_run("API DocDigital — Comunicación de prueba para integración\n")
    run2.font.size = Pt(14)
    run3 = title.add_run("POST /pruebas-integracion/comunicaciones/test-recepcion\n\n")
    run3.font.size = Pt(12)
    run4 = title.add_run("Tarjeta QA-5687 — Manual de Integración API MW v3.5")
    run4.italic = True

    doc.add_paragraph()
    add_kv_table(doc, [
        ("Proyecto", "DD-AUTOMATIZACION-QA — Framework automatización API Doc. Digital"),
        ("Módulo probado", "Pruebas de integración / test-recepcion"),
        ("Ambiente", "QA — https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Versión API", "Middleware v3.5"),
        ("Fecha de ejecución", "24 de junio de 2026"),
        ("Ejecutado por", "Juan Francisco Candia (JFC) — QA Automatizador"),
        ("Repositorio", "JFCandia-Digital/DD-AUTOATIZACION-QA (rama main)"),
        ("Resultado global", "9 casos PASS automatizados | Suite Cucumber: 10 scenarios (10 passed)"),
    ])

    doc.add_page_break()

    add_heading(doc, "1. Objetivo", 1)
    doc.add_paragraph(
        "Documentar las evidencias de las pruebas ejecutadas sobre el endpoint "
        "POST /pruebas-integracion/comunicaciones/test-recepcion, que permite al integrador "
        "enviar comunicaciones oficiales de prueba a su propia entidad para validar métodos "
        "de recepción y rechazo, según criterios de aceptación de la tarjeta QA-5687."
    )

    add_heading(doc, "2. Alcance", 1)
    for item in [
        "Happy path: envío con materia y folio (parámetros de entrada).",
        "Flags opcionales: incorporaAnexos, isReservado, asociarProcedimientoAdministrativo.",
        "Validación de campos obligatorios (materia, folio) — HTTP 400.",
        "Seguridad: token inválido, nulo y expirado — HTTP 401.",
        "Creación individual (un POST genera una comunicación).",
        "Evidencia E2E mediante correos automáticos DocDigital (ids 89721–89726).",
        "Flujo OP receptora: POST → GET pendientes-recepcion → PUT acuse/rechazo (TC-INT-017/018).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fuera de alcance en esta iteración: ").bold = True
    p.add_run(
        "TC-INT-010 a 015 (validaciones GET detalle/archivo/trazabilidad)."
    )

    add_heading(doc, "3. Ambiente y configuración de prueba", 1)
    add_kv_table(doc, [
        ("URL Base", "https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Endpoint", "POST /pruebas-integracion/comunicaciones/test-recepcion"),
        ("Autenticación", "OAuth2 — Token PDI (CLIENT_ID_PDI / CLIENT_SECRET_PDI)"),
        ("Config servidor (Pablo)", "Credenciales API entidad test en INFRA QA (24-06-2026)"),
        ("Documento principal", "Fijo sistema — prueba_dipres_test.pdf (FEA)"),
        ("Despachadora interna", "Agosto KE (login interno servidor)"),
        ("Destinataria", "Entidad del token OAuth"),
        ("Naming QA", "materia: QA Test Recepcion JFC … | folio: QA-REC-JFC-…"),
        ("Framework", "Cucumber + TypeScript + Axios"),
        ("Comando suite POST", "npm run test_recepcion_report"),
        ("Comando suite E2E OP", "npm run test_recepcion_e2e_report"),
        ("Comando cierre QA-5687", "npm run test_recepcion_cierre_report"),
    ])

    add_heading(doc, "4. Historial de bloqueo y destrabe", 1)
    t0 = doc.add_table(rows=1, cols=3)
    t0.style = "Table Grid"
    for i, h in enumerate(["Fecha", "Evento", "Estado"]):
        t0.rows[0].cells[i].text = h
        set_cell_shading(t0.rows[0].cells[i], "1F4E79")
    for row in BLOQUEO:
        r = t0.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v

    add_heading(doc, "5. Resumen ejecutivo de resultados", 1)
    add_kv_table(doc, [
        ("Total casos matriz", "18"),
        ("Casos PASS (automatizados)", "9 (TC-INT-001 a 009 + 016)"),
        ("Casos pendientes", "8 (TC-INT-010 a 015, 017, 018)"),
        ("Escenarios Cucumber @TestRecepcion", "10"),
        ("Steps ejecutados", "70"),
        ("Tiempo ejecución suite", "~14 segundos"),
        ("Comunicaciones creadas", "89721 (probe), 89722, 89723, 89724, 89725, 89726"),
        ("Correos DocDigital recibidos", "6 (89721–89726)"),
        ("Reporte HTML", "reports/index.html — 100% scenarios passed"),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "Tras la configuración de credenciales entidad test por infraestructura (Pablo), "
        "el endpoint respondió HTTP 200 de forma consistente. Los despachos exitosos "
        "generaron correos de recepción DocDigital, confirmando el flujo end-to-end."
    )

    doc.add_page_break()

    add_heading(doc, "6. Matriz de casos de prueba", 1)
    headers = ["ID", "Tipo", "Descripción", "Pasos", "Esperado", "Obtenido", "Estado"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(7)

    for caso in CASOS:
        row = table.add_row().cells
        for i, val in enumerate([caso[0], caso[1], caso[2], caso[3], caso[4], caso[5], caso[6]]):
            row[i].text = val
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)
        if caso[6] == "PASS":
            set_cell_shading(row[6], "C6EFCE")
        elif caso[6] == "Pendiente":
            set_cell_shading(row[6], "FFEB9C")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Matriz completa en Excel: ").bold = True
    p.add_run("docs/api-v3.5/matriz_api_test_recepcion_integracion_DocDigital.xlsx")

    add_heading(doc, "7. Evidencias adjuntas (insertar capturas)", 1)
    evidencias = [
        ("E-01", "Reporte HTML Cucumber", "reports/index.html — 10 scenarios, 100% passed", "Captura dashboard + detalle feature"),
        ("E-02", "Terminal ejecución", "npm run test_recepcion_report — 10 scenarios (10 passed)", "Captura terminal"),
        ("E-03", "Probe manual", "npm run probe:test-recepcion — id 89721", "Captura terminal"),
        ("E-04", "Correos DocDigital", "Comunicaciones 89721, 89722, 89723, 89724, 89725, 89726", "Captura bandeja Gmail no-reply"),
        ("E-05", "Plataforma DocDigital", "Detalle comunicación en Oficina de Partes QA", "Captura web docv3.test"),
        ("E-06", "Matriz de pruebas", "matriz_api_test_recepcion_integracion_DocDigital.xlsx", "Archivo Excel 18 casos TC-INT"),
        ("E-07", "Bloqueo previo", "Swagger 401 ENTIDAD_TEST (23-06) + config Pablo (24-06)", "Capturas Asana / Swagger"),
        ("E-08", "Código automatización", "postTestRecepcion.feature + steps + requestBodies", "Repositorio main"),
    ]
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["ID", "Tipo evidencia", "Descripción", "Formato"]):
        t2.rows[0].cells[i].text = h
        set_cell_shading(t2.rows[0].cells[i], "1F4E79")
    for ev in evidencias:
        r = t2.add_row().cells
        for i, v in enumerate(ev):
            r[i].text = v

    doc.add_paragraph()
    doc.add_paragraph()

    # Espacios para pegar capturas
    for label in [
        "E-01 — Reporte Cucumber (dashboard 10/10):",
        "E-02 — Terminal 10 scenarios passed:",
        "E-03 — Probe test-recepcion id 89721:",
        "E-04 — Bandeja correos 89721–89726:",
        "E-05 — Pantalla DocDigital (detalle comunicación):",
    ]:
        doc.add_paragraph(label).runs[0].bold = True
        doc.add_paragraph("[ Insertar captura aquí ]").runs[0].italic = True
        doc.add_paragraph()

    doc.add_page_break()

    add_heading(doc, "8. Conclusiones", 1)
    doc.add_paragraph(
        "Las pruebas automatizadas del endpoint POST /pruebas-integracion/comunicaciones/test-recepcion "
        "fueron ejecutadas exitosamente en ambiente QA tras la configuración de credenciales entidad test. "
        "Se validaron escenarios positivos con flags opcionales, validaciones de campos obligatorios "
        "y criterios de seguridad OAuth."
    )
    doc.add_paragraph(
        "La suite Cucumber @TestRecepcion reportó 10 escenarios passed (70 steps). "
        "Seis correos DocDigital (ids 89721 a 89726) confirman el procesamiento end-to-end. "
        "Los casos TC-INT-010 a 018 quedan documentados como pendientes para una segunda iteración."
    )
    doc.add_paragraph(
        "Se recomienda el cierre de la tarjeta QA-5687 con entrega QA, dejando los casos "
        "pendientes de GET detalle y E2E acuse/rechazo para seguimiento futuro."
    )

    add_heading(doc, "9. Anexos", 1)
    for item in [
        "Anexo A: matriz_api_test_recepcion_integracion_DocDigital.xlsx",
        "Anexo B: reports/index.html (reporte Cucumber)",
        "Anexo C: Capturas correos DocDigital 89721–89726",
        "Anexo D: postTestRecepcion.feature — src/api-test/apiComunicaciones/features/",
        "Anexo E: Manual API MW v3.5 — docs/api-v3.5/",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.add_run("\n\n_______________________________\n")
    firma.add_run("Juan Francisco Candia (JFC)\n").bold = True
    firma.add_run("QA Automatizador\n")
    firma.add_run("Fecha: 24-06-2026")

    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
