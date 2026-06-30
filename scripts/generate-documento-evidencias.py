#!/usr/bin/env python3
"""Genera Documento_Evidencias_Pruebas_Notificacion.docx"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/api-v3.5/Documento_Evidencias_Pruebas_Notificacion.docx"

CASOS = [
    ("TC-NOTIF-001", "API / Integración", "Validar creación de comunicación tipo Notificación con documento firmado y datos PA válidos",
     "POST /comunicaciones/despachar-tipo-notificacion | Validar 200 + result.id",
     "200 OK, comunicación creada", "200 OK — id 89685 (23-06-2026 10:38:08)", "PASS",
     "reports/index.html; correo DocDigital 89685"),
    ("TC-NOTIF-002", "API / Validación", "Validar error sin procedimiento administrativo",
     "POST sin tipoProcedimientoAdministrativo", "400 / errorCode 40000", "400 — tipoProcedimientoAdministrativo es obligatorio", "PASS", "reports/index.html"),
    ("TC-NOTIF-003", "API / End-to-End", "Validar flujo notificación y recepción en DocDigital",
     "POST exitoso + verificar correo no-reply DocDigital", "POST 200 + correo recepción",
     "POST 200 — ids 89685, 89687, 89688; correos: 89679, 89682, 89685, 89687", "PASS",
     "Captura correos; reports/index.html"),
    ("TC-NOTIF-004", "API / Validación", "Rechazo documento sin firma digital",
     "POST con Prueba.pdf sin FEA", "400 / errorCode 4001", "400 — errorCode 4001", "PASS", "reports/index.html"),
    ("TC-NOTIF-005", "API / Validación", "Error sin destinatarios", "POST sin destinatarios", "400 / 40000",
     "400 — Se requiere al menos una entidad destinataria", "PASS", "reports/index.html"),
    ("TC-NOTIF-006", "API / Validación", "Error sin configuracionDestinatarios", "POST sin configuración", "400 / 40000",
     "400 — configuracionDestinatarios es obligatorio", "PASS", "reports/index.html"),
    ("TC-NOTIF-007", "API / Validación", "Error sin usuario solicitante", "POST sin usuarioSolicitante", "400 / 40000",
     "400 — Usuario solicitante es obligatorio", "PASS", "reports/index.html"),
    ("TC-NOTIF-008", "API / Validación", "Error sin RUN usuario", "POST sin run", "400 / 40000",
     "400 — RUN del usuario es obligatorio", "PASS", "reports/index.html"),
    ("TC-NOTIF-009", "API / Validación", "Error sin DV usuario", "POST sin dv", "400 / 40000",
     "400 — Digito verificador del RUN es obligatorio", "PASS", "reports/index.html"),
    ("TC-NOTIF-010", "API / Negocio", "Dependiente 4758 en copia", "POST JSON_NOTIFICACION_DEPENDIENTE_EN_COPIA", "200 + id",
     "200 — id 89687 + correo DocDigital", "PASS", "reports/index.html; correo 89687"),
    ("TC-NOTIF-011", "API / Negocio", "Dependiente 4758 como principal", "POST JSON_NOTIFICACION_DEPENDIENTE_COMO_PRINCIPAL",
     "400 (manual) / 200 (QA)", "200 — id 89688", "PASS", "Hallazgo H-01 documentado"),
    ("TC-NOTIF-012", "API / Negocio", "Destinatario no dependiente", "POST JSON_NOTIFICACION_DESTINATARIO_NO_DEPENDIENTE",
     "400 rechazo negocio", "400 — errorCode 4001", "PASS", "Hallazgo H-02: cobertura parcial"),
]

HALLAZGOS = [
    ("H-01", "Media", "Manual indicaba rechazo de dependiente como principal; API QA acepta despacho 598→4758 con 200",
     "Confirmar regla de negocio con equipo backend/funcional"),
    ("H-02", "Baja", "Caso destinatario no dependiente valida error 4001 (firma) antes que regla de dependiente",
     "Profundizar en iteración futura con PDF firmado"),
]


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    return table


def main():
    doc = Document()

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENTO DE EVIDENCIAS DE PRUEBAS\n")
    run.bold = True
    run.font.size = Pt(16)
    run2 = title.add_run("API DocDigital — Endpoint Notificación\n")
    run2.font.size = Pt(14)
    run3 = title.add_run("POST /comunicaciones/despachar-tipo-notificacion\n\n")
    run3.font.size = Pt(12)
    run4 = title.add_run("Manual de Integración API MW v3.5")
    run4.italic = True

    doc.add_paragraph()
    add_kv_table(doc, [
        ("Proyecto", "DD-AUTOMATIZACION-QA — Framework automatización API Doc. Digital"),
        ("Módulo probado", "Comunicaciones / Notificación"),
        ("Ambiente", "QA — https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Versión API", "Middleware v3.5"),
        ("Fecha de ejecución", "23 de junio de 2026"),
        ("Ejecutado por", "Juan Candia — QA Automatizador"),
        ("Repositorio", "JFCandia-Digital/DD-AUTOMATIZACION-QA"),
        ("Resultado global", "12 casos de prueba — 12 PASS | Suite Cucumber: 11 scenarios (11 passed)"),
    ])

    doc.add_page_break()

    # 1. Objetivo
    add_heading(doc, "1. Objetivo", 1)
    doc.add_paragraph(
        "Documentar las evidencias de las pruebas automatizadas ejecutadas sobre el endpoint "
        "POST /comunicaciones/despachar-tipo-notificacion, en el marco de la validación QA "
        "solicitada en la tarjeta Asana, alineada al Manual de Integración API MW v3.5."
    )

    # 2. Alcance
    add_heading(doc, "2. Alcance", 1)
    for item in [
        "Flujo exitoso de despacho tipo notificación (documento con firma electrónica válida).",
        "Validación de rechazo por documento sin firma (errorCode 4001).",
        "Validación de campos obligatorios del Procedimiento Administrativo (errorCode 40000).",
        "Reglas de negocio con entidades dependientes (copia, principal, no dependiente).",
        "Evidencia E2E mediante correos automáticos de DocDigital.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Fuera de alcance en esta iteración: ").bold = True
    p.add_run("automatización de GET pendientes-recepción y PUT acuse-recibo.")

    # 3. Ambiente
    add_heading(doc, "3. Ambiente y configuración de prueba", 1)
    add_kv_table(doc, [
        ("URL Base", "https://middleware.docv3.test.digital.gob.cl/api/v3"),
        ("Autenticación", "OAuth2 — Token PDI (CLIENT_ID_PDI / CLIENT_SECRET_PDI)"),
        ("Entidad despachadora", "598"),
        ("Destinatario principal", "156"),
        ("Entidad dependiente QA", "4758 (PI Inspectoría General — GET /entidades/dependientes/598)"),
        ("PDF con firma", "2_FIRMANTES_EN_DOC_DIGITAL.pdf"),
        ("PDF sin firma", "Prueba.pdf"),
        ("Framework", "Cucumber + TypeScript + Axios"),
        ("Comando suite", "npm run notificacion_report"),
    ])

    # 4. Resumen ejecutivo
    add_heading(doc, "4. Resumen ejecutivo de resultados", 1)
    add_kv_table(doc, [
        ("Total casos de prueba (matriz)", "12"),
        ("Casos exitosos (PASS)", "12"),
        ("Casos fallidos (FAIL)", "0"),
        ("Escenarios Cucumber @Notificacion", "11"),
        ("Steps ejecutados", "97"),
        ("Tiempo ejecución suite", "~14 segundos"),
        ("Comunicaciones creadas (ejecución final)", "89685, 89687, 89688"),
        ("Reporte HTML", "reports/index.html — 100% scenarios passed"),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "Se validó el comportamiento real de la API en ambiente QA. Los despachos exitosos "
        "generaron correos de notificación DocDigital (remitente no-reply), confirmando el "
        "procesamiento end-to-end del flujo de notificación."
    )

    doc.add_page_break()

    # 5. Matriz de casos
    add_heading(doc, "5. Matriz de casos de prueba ejecutados", 1)
    headers = ["ID", "Tipo", "Descripción", "Pasos", "Esperado", "Obtenido", "Estado"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(8)

    for caso in CASOS:
        row = table.add_row().cells
        for i, val in enumerate([caso[0], caso[1], caso[2], caso[3], caso[4], caso[5], caso[6]]):
            row[i].text = val
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
        if caso[6] == "PASS":
            set_cell_shading(row[6], "C6EFCE")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Matriz completa en Excel: ").bold = True
    p.add_run("docs/api-v3.5/matriz_api_notificacion_DocDigital.xlsx")

    # 6. Evidencias
    add_heading(doc, "6. Evidencias adjuntas", 1)
    evidencias = [
        ("E-01", "Reporte HTML Cucumber", "reports/index.html — 11 scenarios, 100% passed", "Captura pantalla reporte"),
        ("E-02", "Terminal ejecución", "npm run notificacion_report — 11 scenarios (11 passed)", "Captura terminal"),
        ("E-03", "Correos DocDigital", "Comunicaciones 89679, 89682, 89685, 89687", "Captura bandeja correo no-reply"),
        ("E-04", "Matriz de pruebas", "matriz_api_notificacion_DocDigital.xlsx", "Archivo Excel 12 casos TC-NOTIF"),
        ("E-05", "Código automatización", "postNotificacion.feature + steps + requestBodies", "Repositorio DD-AUTOMATIZACION-QA"),
        ("E-06", "Respuestas API", "JSON 200/400 con errorCode 4001 y 40000", "Attachments en reporte Cucumber"),
    ]
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(["ID", "Tipo evidencia", "Descripción", "Formato"]):
        t2.rows[0].cells[i].text = h
        set_cell_shading(t2.rows[0].cells[i], "1F4E79")
    for ev in evidencias:
        r = t2.add_row().cells
        for i, v in enumerate(ev):
            r[i].text = v

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[ Insertar aquí capturas: reporte HTML, terminal, correos DocDigital ]").italic = True

    doc.add_page_break()

    # 7. Hallazgos
    add_heading(doc, "7. Hallazgos e incidencias", 1)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    for i, h in enumerate(["ID", "Severidad", "Descripción", "Acción recomendada"]):
        t3.rows[0].cells[i].text = h
        set_cell_shading(t3.rows[0].cells[i], "1F4E79")
    for h in HALLAZGOS:
        r = t3.add_row().cells
        for i, v in enumerate(h):
            r[i].text = v

    # 8. Conclusiones
    add_heading(doc, "8. Conclusiones", 1)
    doc.add_paragraph(
        "Las pruebas automatizadas del endpoint POST /comunicaciones/despachar-tipo-notificacion "
        "fueron ejecutadas exitosamente en ambiente QA. Se validaron escenarios positivos, "
        "negativos de validación de campos PA y reglas de negocio con entidades dependientes."
    )
    doc.add_paragraph(
        "La suite Cucumber @Notificacion reportó 11 escenarios passed (97 steps). "
        "La evidencia E2E mediante correos DocDigital confirma que las comunicaciones fueron "
        "procesadas correctamente por el sistema."
    )
    doc.add_paragraph(
        "Se recomienda el cierre de la tarjeta Asana, dejando registrados los hallazgos H-01 y H-02 "
        "para seguimiento con el equipo funcional."
    )

    # 9. Anexos
    add_heading(doc, "9. Anexos", 1)
    for item in [
        "Anexo A: matriz_api_notificacion_DocDigital.xlsx",
        "Anexo B: reports/index.html (reporte Cucumber)",
        "Anexo C: Capturas correos DocDigital",
        "Anexo D: Manual API MW v3.5 — docs/api-v3.5/",
        "Anexo E: Código — src/api-test/apiComunicaciones/features/postNotificacion.feature",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Firma
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.add_run("\n\n_______________________________\n").bold = False
    firma.add_run("Juan Candia\n").bold = True
    firma.add_run("QA Automatizador\n")
    firma.add_run("Fecha: 23-06-2026")

    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    main()
